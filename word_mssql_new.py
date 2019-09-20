# -*- coding:utf-8 -*-
import pymssql  
import sys
import codecs
import os,time
import docx
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import *
import openpyxl
from openpyxl import load_workbook
from openpyxl.styles import NamedStyle
import concurrent.futures
import printer_api as printer
import configparser 

from copy import deepcopy

config = configparser .RawConfigParser()
config.read('word_config.ini',encoding='utf-8')


nn=0
time_sum=0
auto_printer_mode = config.get('word', 'auto_printer_mode')
inputsqlcode_mode = config.get('word', 'inputsqlcode_mode')
sql_code_mode = config.get('word', 'sql_code_mode')

print ('auto_printer_mode:',auto_printer_mode)
print ('inputsqlcode_mode:',inputsqlcode_mode)

sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

if os.path.isdir('bill'):

    if os.path.isfile('bill/finish.txt'):
        os.remove('bill/finish.txt')
    if os.path.isfile('bill/pass.txt'):
        os.remove('bill/pass.txt')
else:
    os.makedirs('bill') 

data=[]
list_date=[]
list_name=[]
list_namecode=[]
list_transID=[]
list_docnum=[]
list_JrnlMemo=[]
list_doccur=[]
list_sellsname=[]
list_code=[]
list_tel=[]
list_address=[]
list_doctotal=[]
list_paidtodate=[]
list_paidsum=[]

error_num=0

def remove_row(table, row):
    try:

        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
    except ValueError:
        print('Error')
        
def encode_decode(ts):
    s = ts
    #print(s.encode("utf8").decode("cp950", "ignore"))
    return s.encode("utf8").decode("cp950", "ignore")


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
def error_mess(ename,tname):
    if ename==tname:
        doc.save('bill/form_Error.docx')
        
def add_dot(m):
    mc=''
    m=m[::-1]
    
    n=1
    for i in m:

        if n%3==0 and n!=0:
            mc+=i+','
        else:
            mc+=i
        n+=1
    return mc[::-1]


conn = pymssql.connect(server=config.get('param', 'server_ip'), user=config.get('param', 'user'), password=config.get('param', 'pass'), database=config.get('param', 'db'))  
cursor = conn.cursor()  

if sql_code_mode==1:
    sqlstr = config.get('param', 'sql_code')
else:
    sqlstr=""
#sqlstr = encode_decode(sqlstr)


if inputsqlcode_mode==1:
    sqlcode=''
    sqlcode=input('enter sql code:')
    cursor.execute(sqlcode)
#print (sqlcode)
else:
    cursor.execute(sqlstr)  

s=time.time()

for row in cursor:
    data.append(row)
    
    list_date.append(row[0])
    
    list_transID.append(row[1])
    list_docnum.append(row[2])
    list_JrnlMemo.append(row[3])
    list_doccur.append(row[4])
    list_sellsname.append(row[5])
    list_code.append(row[6])
    list_tel.append(row[9])
    list_address.append(row[10])
    list_doctotal.append(row[11])
    list_paidtodate.append(row[12])
    list_paidsum.append(row[13])
    list_name.append(row[7])
    list_namecode.append(row[8])
    #print (row)

conn.rollback()  
conn.close()    
e=time.time()
print ('Load SQL Time: ',e-s)
#print (data[1][0])
#print (len(data))

set_name=set(list_name)


with concurrent.futures.ProcessPoolExecutor() as executor:

    for name in set_name:
        docsum=0
        #if list_name.count(name)<=500 :
            
        start=time.time()

        doc = Document('form_final_new_top.docx')
        doc_copy = Document('form_final_new.docx')
        table=doc.tables[1]
        remove_row(table,table.rows[0])


        print(name)
        
        
        try:
            num=0
            for i in range(0,len(data)):
                #num=num+1
                if list_name[i]==name:
                    error_num=i
                    
                    
                    
                    
                    if num%35==0 and num!=0:

                        pn=0

                        
                        for p in doc.paragraphs:
                            
                            if pn==len(doc.paragraphs)-3:
                                
                                run = p.add_run()
                                run.add_break(WD_BREAK.PAGE)
                            
                            if pn==len(doc.paragraphs)-2:
                                
                                p.text='客戶名稱 '+name
                                paragraph_format = p.paragraph_format
                                paragraph_format.left_indent = Pt(-20)
                                run_obj = p.runs
                                run = run_obj[0]
                                font = run.font
                                font.size = Pt(10)
                            if pn==len(doc.paragraphs)-1:
                                p.text='客戶代號 '+namecode
                                paragraph_format = p.paragraph_format
                                paragraph_format.left_indent = Pt(-20)
                                run_obj = p.runs
                                run = run_obj[0]
                                font = run.font
                                font.size = Pt(10)

                            pn+=1
                    
                        copy_table_after(doc_copy.tables[0],doc.paragraphs[len(doc.paragraphs)-1])

                        doc.add_paragraph('')
                        
                        copy_table_after(doc_copy.tables[1],doc.paragraphs[len(doc.paragraphs)-1])
                        remove_row(doc.tables[int(num/35)+int(num/35)+1],doc.tables[int(num/35)+int(num/35)+1].rows[0])
                        

                        doc.add_paragraph('')
                        doc.add_paragraph('')
                        doc.add_paragraph('')
                    #else:
                    try:
                    
                        new_row=doc.tables[int(num/35)+int(num/35)+1].add_row().cells
                        
                        new_row[0].text = str(list_date[i])[0:10]
                        
                        new_row[2].text = str(list_transID[i])
                        new_row[3].text = str(list_docnum[i])
                        
                        new_row[5].text = str(int(list_paidtodate[i]))
                        new_row[6].text = str(int(list_paidsum[i]))
                        new_row[7].text = str(list_doccur[i])
                        #new_row[8].text = ''#str(ws['H'+str(num)].value)
                        #new_row[9].text = ''#str(ws['H'+str(num)].value)
                        new_row[8].text = str(list_sellsname[i])
                        new_row[9].text = str(list_code[i])
                        #print(str(list_JrnlMemo[i]))
                        if "-" in str(list_JrnlMemo[i]):
                            new_row[1].text = str(list_JrnlMemo[i]).split('-')[0]
                            if "," not in str(list_JrnlMemo[i]):
                                new_row[4].text = ''
                            else:
                                new_row[4].text = str(list_JrnlMemo[i]).split('-')[1].split(',')[1]
                            
                        elif "," in str(list_JrnlMemo[i]):
                            new_row[1].text = str(list_JrnlMemo[i]).split(',')[0]
                            new_row[4].text = str(list_JrnlMemo[i]).split(',')[1]
                            
                        elif ";" in str(list_JrnlMemo[i]):
                            new_row[1].text = str(list_JrnlMemo[i]).split(';')[0]
                            new_row[4].text = str(list_JrnlMemo[i]).split(';')[1]

                        else:
                            new_row[1].text = str(list_JrnlMemo[i]).split('-')[0]
                            new_row[4].text = ''
                        
                    
                        docsum+=int(list_doctotal[i])
                        
                        namecode=str(list_namecode[i])
                        tel=str(list_tel[i])
                        if '\r' in str(list_address[i]):
                            list_address[i]=str(list_address[i]).replace('\r', '')
                        addres=str(list_address[i])
                #else:
                    
                        for row in doc.tables[int(num/35)+int(num/35)+1].rows:
                            for cell in row.cells:
                                paragraphs = cell.paragraphs
                                paragraph = paragraphs[0]
                                run_obj = paragraph.runs
                                run = run_obj[0]
                                font = run.font
                                font.size = Pt(10)
                    
                    except ValueError:
                        error_mess(name,name)
                        print('Error')
            #print(nn)
                    num+=1
                    
            n=0
            for para in doc.paragraphs :

                if n==10:
                    #print (n)
                    para.text=name+' 會計部 收'
                    #print(para.text)
                elif n==34:
                    #print (n)
                    para.text=name
                elif n==12 or n==36:
                    #print (n)
                    para.text=addres
                elif n==14 or n==38:
                    #print (n)
                    para.text=tel
                elif n==43:
                    #print (n)
                    para.text='客戶名稱 '+name
                    paragraph_format = p.paragraph_format
                    paragraph_format.left_indent = Pt(-20)
                    run_obj = para.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(10)
                elif n==44:
                    #print (n)
                    para.text='客戶代號 '+namecode
                    paragraph_format = p.paragraph_format
                    paragraph_format.left_indent = Pt(-20)
                    run_obj = p.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(0)
                n+=1
            
            if int(35-(num%35))<=5:
                pn=0
                for p in doc.paragraphs:
                    if pn==len(doc.paragraphs)-3:
                        
                        run = p.add_run()
                        run.add_break(WD_BREAK.PAGE)
                        
                    pn+=1
            
            copy_table_after(doc_copy.tables[2],doc.paragraphs[len(doc.paragraphs)-3])
            


            
            table_buttom=doc.tables[len(doc.tables)-1]

            row1=table_buttom.rows[0].cells
            row2=table_buttom.rows[1].cells
            row3=table_buttom.rows[2].cells
            
            row1[2].text='TWD'
            row2[3].text=add_dot(str(docsum))

            row3[3].text=add_dot(str(docsum))
            row2[4].text=str(num)
            
            
            for row in table_buttom.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(10)

            delete_paragraph(doc.paragraphs[len(doc.paragraphs)-1])
            delete_paragraph(doc.paragraphs[len(doc.paragraphs)-1])
            
            doc.save('bill/form_'+name+'.docx')
            
            end=time.time()

            if auto_printer_mode==1:
                printer.printer_run('bill/form_'+name+'.docx')

            
            
        
            time_sum+=(end-start)
            print ("Individual Time: ",end-start)
    #print (name)
            with open('bill/finish.txt','a') as f:
                f.write(name)
                f.write('\n')
        except ValueError:
            print('Error :',name)
            with open('bill/finish.txt','a') as f:
                f.write(list_namecode[error_num])
                f.write(",")
                f.write(str(list_docnum[error_num]))
                f.write(", Error")
                f.write('\n')
            
        #else:
        #    with open('bill/pass.txt','a') as f:
        #        f.write(name)
        #        f.write('\n')
            #print('Pass:',encode_decode(name))
        #    print ('Pass:',name)
print ("Total Spend Time: " ,time_sum)
print ('Done!!')
